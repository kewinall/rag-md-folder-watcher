from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document

from .base import BaseConverter
from .ocr import UnsupportedImageForOcr, ocr_available, ocr_image_bytes
from ..utils import table_to_markdown


class DocxConverter(BaseConverter):
    extensions = {".docx"}
    name = "docx-with-image-ocr"

    def convert(self, path: Path) -> str:
        document = Document(str(path))
        lines: list[str] = [f"# {path.stem}", ""]

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name or "").lower() if paragraph.style else ""
            if "heading 1" in style or style == "title":
                lines.append(f"# {text}")
            elif "heading 2" in style:
                lines.append(f"## {text}")
            elif "heading 3" in style:
                lines.append(f"### {text}")
            else:
                lines.append(text)
            lines.append("")

        for table_index, table in enumerate(document.tables, 1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            markdown_table = table_to_markdown(rows)
            if markdown_table:
                lines.extend([f"## 表格 {table_index}", markdown_table])

        if self.config.ocr and ocr_available():
            self._append_image_ocr(document, lines)

        return "\n".join(lines) + "\n"

    def _append_image_ocr(self, document: Document, lines: list[str]) -> None:
        image_parts: list[tuple[str, str, bytes]] = []
        seen_hashes: set[str] = set()

        # package.parts covers images used by the body, headers and footers.
        for part in document.part.package.parts:
            content_type = str(getattr(part, "content_type", "") or "")
            if not content_type.lower().startswith("image/"):
                continue

            blob = bytes(getattr(part, "blob", b"") or b"")
            if not blob:
                continue

            digest = hashlib.sha256(blob).hexdigest()
            if digest in seen_hashes:
                continue
            seen_hashes.add(digest)

            image_parts.append((str(part.partname), content_type, blob))

        recognized: list[tuple[int, str, str]] = []
        unsupported_details: list[str] = []
        failed_details: list[str] = []
        unsupported_count = 0
        failed_count = 0
        empty_count = 0

        for image_index, (part_name, content_type, blob) in enumerate(image_parts, 1):
            try:
                text = ocr_image_bytes(
                    blob,
                    self.config,
                    source_name=part_name,
                    content_type=content_type,
                )
            except UnsupportedImageForOcr as exc:
                unsupported_count += 1
                if len(unsupported_details) < 10:
                    unsupported_details.append(
                        f"圖片 {image_index}（{part_name}，{content_type or '未知類型'}）：{exc}"
                    )
                continue
            except Exception as exc:
                failed_count += 1
                if len(failed_details) < 10:
                    failed_details.append(
                        f"圖片 {image_index}（{part_name}，{content_type or '未知類型'}）："
                        f"{type(exc).__name__}: {exc}"
                    )
                continue

            if text:
                recognized.append((image_index, part_name, text))
            else:
                empty_count += 1

        if recognized:
            lines.append("## 文件圖片 OCR")
            for image_index, part_name, text in recognized:
                lines.extend([
                    f"### 圖片 {image_index}",
                    f"<!-- source-part: {part_name} -->",
                    text,
                    "",
                ])

        if image_parts:
            self.warnings.append(
                "Word 圖片 OCR 摘要："
                f"圖片 {len(image_parts)} 張、辨識到文字 {len(recognized)} 張、"
                f"無文字 {empty_count} 張、略過不支援 {unsupported_count} 張、"
                f"其他失敗 {failed_count} 張"
            )

        if unsupported_details:
            self.warnings.append(
                "Word 圖片略過明細（最多 10 筆）：" + " | ".join(unsupported_details)
            )
        if failed_details:
            self.warnings.append(
                "Word 圖片 OCR 失敗明細（最多 10 筆）：" + " | ".join(failed_details)
            )
